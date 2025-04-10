from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

class DocumentGenerator:
    """Service for generating resume documents in DOCX and PDF formats"""
    
    def create_resume_document(self, resume_data, filename):
        """Generate a resume document from structured data and save to a file"""
        doc = self._generate_document(resume_data)
        
        # Save the document to a file
        doc.save(filename)
        print(f"Resume saved as {filename}")
        return filename
    
    def create_resume_document_stream(self, resume_data):
        """Generate a resume document and return as in-memory BytesIO stream"""
        doc = self._generate_document(resume_data)
        
        # Save to an in-memory file-like object
        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0) # Reset pointer to beginning of stream
        
        return output_stream
    
    def create_pdf_stream(self, resume_data):
        """Generate a PDF document in memory
        
        Returns:
            BytesIO: PDF content as a stream or None if conversion failed
        """
        # Create the DOCX in memory first
        docx_stream = self.create_resume_document_stream(resume_data)
        
        try:
            from docx2pdf import convert
            import tempfile
            import os
            
            # We need to save to temp files for the conversion
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                temp_docx.write(docx_stream.getvalue())
                temp_docx_path = temp_docx.name
            
            temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
            
            # Convert docx to pdf
            convert(temp_docx_path, temp_pdf_path)
            
            # Read the PDF into memory
            with open(temp_pdf_path, 'rb') as f:
                pdf_bytes = f.read()
            
            # Clean up temp files
            try:
                os.remove(temp_docx_path)
                os.remove(temp_pdf_path)
            except Exception as e:
                print(f"Error cleaning up temp files: {str(e)}")
            
            # Return as BytesIO
            pdf_stream = BytesIO(pdf_bytes)
            return pdf_stream
            
        except ImportError:
            print("docx2pdf module not installed")
            return None
        except Exception as e:
            print(f"PDF conversion error: {str(e)}")
            return None
    
    def _generate_document(self, resume_data):
        """Core document generation functionality shared by file and stream methods"""
        if isinstance(resume_data, str):
            try:
                resume_data = json.loads(resume_data)
            except json.JSONDecodeError as e:
                raise ValueError(f"Invalid JSON string: {str(e)}") from e
            
        if 'contact' not in resume_data or not isinstance(resume_data['contact'], dict):
            raise ValueError("Missing or invalid contact information")
        
        # Create a new Document
        doc = Document()
        
        # ===== Set Default Styles =====
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Helvetica Neue'
        font.size = Pt(10)

        # Set 0.5 inch margins on all sides
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # ===== Contact Information =====
        # More compact header with name and contact info on same line
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
        contact.add_run(resume_data['contact']['name']).bold = True
        
        # Contact info on same line with minimal spacing
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(0)  # No space after
        contact_p.paragraph_format.space_before = Pt(0)  # No space before
        
        # Add email with hyperlink
        email = resume_data['contact']['email'].strip()
        self._add_hyperlink(contact_p, email, f"mailto:{email}")
        contact_p.add_run(" || ").bold = True
        
        # Add phone number
        contact_p.add_run(resume_data['contact']['phone'].strip()).bold = True
        contact_p.add_run(" || ").bold = True
        
        # Add LinkedIn with hyperlink - keep exactly as provided in the input
        linkedin_text = resume_data['contact']['linkedin'].strip()
        linkedin_url = linkedin_text
        
        # Ensure URL has https:// prefix for proper linking if not already present
        if not linkedin_url.startswith("http"):
            linkedin_url = "https://" + linkedin_url
        
        self._add_hyperlink(contact_p, linkedin_text, linkedin_url)
        
        # Add title on same line or with minimal spacing
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(0)  # No space after
        title_p.paragraph_format.space_before = Pt(2)  # Minimal space before
        title_run = title_p.add_run(resume_data['title'])
        title_run.underline = True
        title_run.bold = True

        # We still need some minimal separation before content sections
        doc.add_paragraph("", style="Normal")
        
        # ===== Professional Summary =====
        summary_heading = doc.add_paragraph()
        summary_heading.add_run('Professional Summary:').bold = True

        summary = doc.add_paragraph(resume_data['summary'])

        # ===== Technical Skills Heading =====
        skills_heading = doc.add_paragraph()
        skills_heading.add_run('Technical Skills:').bold = True

        # Create table with 1 row (to define column widths)
        skills_table = doc.add_table(rows=1, cols=2)
        skills_table.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align with document margins
        skills_table.autofit = False  # Prevent auto-adjustment
        self._set_table_borders(skills_table)

        # Set widths for each column
        skills_table.columns[0].width = Inches(2)
        skills_table.columns[1].width = Inches(5.5)

        # Remove the first placeholder row
        skills_table._tbl.remove(skills_table.rows[0]._tr)

        # Set header texts
        row = skills_table.add_row().cells
        row[0].text = "Category"
        row[1].text = "Skills/Tools"

        # Make both headers bold
        row[0].paragraphs[0].runs[0].bold = True
        row[1].paragraphs[0].runs[0].bold = True

        # Remove spacing after paragraph
        for cell in row:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = 0

        # Add actual data rows - preserve the exact order from input data
        if isinstance(resume_data['technical_skills'], dict):
            # Use skill_order if available, otherwise fall back to dict keys
            skill_order = resume_data.get('skill_order', list(resume_data['technical_skills'].keys()))
            
            # We need to iterate through items in the exact order they appear in the original data
            for category in skill_order:
                if category in resume_data['technical_skills']:
                    skills = resume_data['technical_skills'][category]
                    row = skills_table.add_row().cells
                    row[0].text = category
                    row[1].text = ", ".join(skills) if isinstance(skills, list) else skills

                    # Bold left column text
                    if row[0].paragraphs[0].runs:
                        row[0].paragraphs[0].runs[0].bold = True

                    # Remove spacing after paragraph
                    for cell in row:
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_after = 0

        self._add_section_separator(doc)

        # ===== Work Experience =====
        work_exp_heading = doc.add_paragraph()
        work_exp_heading.add_run('Work Experience').bold = True

        for job in resume_data['work_experience']:
            self._add_work_experience(doc, job)
            
            # Achievements
            for achievement in job['achievements']:
                achievement = achievement.replace("●", "").strip()
                doc.add_paragraph(achievement, style='ListBullet')

        # ===== Education =====
        edu_heading = doc.add_paragraph()
        edu_heading.add_run('Education:').bold = True
        self._add_education(doc, resume_data['education'])
        
        return doc

    def _set_table_borders(self, table):
        """Set borders for a table"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # thin line
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def _add_work_experience(self, doc, job):
        """Add work experience section to document"""
        # Create a 1-row, 2-column table
        table = doc.add_table(rows=2, cols=2)

        # Left cell (Company and Location)
        left_top = table.cell(0, 0)
        left_p = left_top.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_p.paragraph_format.space_after = Pt(0)  # Remove space after
        left_run_company = left_p.add_run(job['company'])
        left_run_company.italic = True
        left_run_company.bold = True

        left_bottom = table.cell(1, 0)
        left_p = left_bottom.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_p.paragraph_format.space_after = Pt(0)  # Remove space after
        left_run_role = left_p.add_run(job['role'])
        left_run_role.italic = True
        
        # Right cell (Position and Dates)
        right_top = table.cell(0, 1)
        right_p = right_top.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_after = Pt(0)  # Remove space after
        right_run_location = right_p.add_run(job['location'])
        right_run_location.bold = True
        right_run_location.italic = True

        right_bottom = table.cell(1, 1)
        right_p = right_bottom.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_after = Pt(0)  # Remove space after
        # Fix date ranges by explicitly replacing Unicode characters
        dates = job['dates'].replace('\u00e2', '-').replace('\u2013', '-').replace('\u2014', '-').replace('  ', ' ')
        right_run_dates = right_p.add_run(dates)
        right_run_dates.italic = True
        
        # Remove table borders
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    def _add_education(self, doc, education):
        """Add education section to document"""
        # Create a 1-row, 2-column table
        table = doc.add_table(rows=2, cols=2)

        # Left cell (institution and degree)
        left_top = table.cell(0, 0)
        left_p = left_top.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_p.paragraph_format.space_after = Pt(0)  # Remove space after
        left_run_institution = left_p.add_run(education['institution'])
        left_run_institution.italic = True
        left_run_institution.bold = True

        left_bottom = table.cell(1, 0)
        left_bottom.width = Inches(5)  # Set width for left cell
        left_p = left_bottom.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_p.paragraph_format.space_after = Pt(0)  # Remove space after
        
        # Fix apostrophe character encoding issues
        degree_text = education['degree']
        degree_text = degree_text.replace('â', "'").replace('â', "'").replace('\u00e2', "'")
        
        left_run_degree = left_p.add_run(degree_text)
        left_run_degree.italic = True
        
        # Right cell (Location and Grad year)
        right_top = table.cell(0, 1)
        right_p = right_top.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_after = Pt(0)  # Remove space after
        right_run_location = right_p.add_run(education['location'])
        right_run_location.bold = True
        right_run_location.italic = True

        right_bottom = table.cell(1, 1)
        right_p = right_bottom.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_after = Pt(0)  # Remove space after
        dates = education['graduation_year'].replace('\u00e2', '-').replace('  ', ' ')
        right_run_dates = right_p.add_run(f"Graduation: {dates}")
        right_run_dates.italic = True

        # Remove table borders
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    def _add_section_separator(self, doc):
        """Add space between sections"""
        doc.add_paragraph()  # Add an empty paragraph for spacing

    def _add_hyperlink(self, paragraph, text, url):
        """Add a hyperlink to a paragraph"""
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Create a run element
        new_run = OxmlElement('w:r')
        
        # Create a bold element if needed
        if paragraph.runs and paragraph.runs[-1].bold:
            b = OxmlElement('w:b')
            new_run.append(b)
        
        # Create a rPr element
        rPr = OxmlElement('w:rPr')
        
        # Add color if you want it
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')  # Blue color
        rPr.append(color)
        
        # Add underline
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        # Create a text element
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        hyperlink.append(new_run)
        
        # Add the hyperlink to the paragraph
        paragraph._p.append(hyperlink)
        
        return hyperlink