from flask import Flask, request, render_template, jsonify, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
from structurer import parse_resume

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/submit', methods=['POST'])
def submit():
    resume_text = request.form.get('resume')
    if not resume_text or not resume_text.strip():
        return "Error: Resume text cannot be empty.", 400

    try:
        parsed_resume = parse_resume(resume_text)
        if parsed_resume:
            # Extract skill categories in their original order
            skill_order = list(parsed_resume['technical_skills'].keys())
            
            # Add the order information to the response
            return jsonify({
                "status": "success", 
                "message": "Resume successfully processed.",
                "data": parsed_resume,
                "skill_order": skill_order  # Include original skill category order
            })
        else:
            return jsonify({"status": "error", "message": "Failed to parse resume."}), 500
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/generate', methods=['POST'])
def generate_document():
    """Generate and download a resume document from the edited data"""
    try:
        # Get the updated resume data from the request
        resume_data = request.json
        
        if not resume_data:
            return jsonify({"status": "error", "message": "No resume data provided"}), 400
            
        # Generate document with a temporary prefix
        temp_docx_filename = f"temp_{resume_data['contact']['name']}.docx"
        
        # Create the document with the temporary name
        doc = Document()
        # All document creation code is in this function
        filename = create_resume_document(resume_data, temp_docx_filename)
        file_path = os.path.join(os.getcwd(), temp_docx_filename)
        
        # Return the file path for download
        return jsonify({
            "status": "success",
            "message": f"Resume successfully generated.",
            "filename": temp_docx_filename,
            "file_path": file_path
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/generate/pdf', methods=['POST'])
def generate_pdf():
    """Generate a PDF version of the resume from the edited data"""
    try:
        # Check if docx2pdf is installed
        try:
            import docx2pdf
        except ImportError:
            return jsonify({
                "status": "error", 
                "message": "PDF conversion requires docx2pdf module. Please install it with 'pip install docx2pdf'. "
                           "Note: docx2pdf requires Microsoft Word or LibreOffice to be installed on your system."
            }), 500
            
        # Get the updated resume data from the request
        resume_data = request.json
        
        if not resume_data:
            return jsonify({"status": "error", "message": "No resume data provided"}), 400
            
        # Generate DOCX document first with a temporary name
        temp_docx_filename = f"temp_{resume_data['contact']['name']}.docx"
        
        # Store original name to use for PDF
        final_name = resume_data['contact']['name']
        
        # Generate temp document
        create_resume_document(resume_data, temp_docx_filename)
        docx_path = os.path.join(os.getcwd(), temp_docx_filename)
        
        # Convert to PDF with the final name
        temp_pdf_filename = f"temp_{final_name}.pdf"
        pdf_path = os.path.join(os.getcwd(), temp_pdf_filename)
        
        # Convert DOCX to PDF using docx2pdf
        try:
            from docx2pdf import convert
            print(f"Converting {docx_path} to {pdf_path}")
            convert(docx_path, pdf_path)
            print("Conversion completed")
            
            # Delete the temporary DOCX file
            if os.path.exists(docx_path):
                os.remove(docx_path)
                print(f"Temporary DOCX file {temp_docx_filename} deleted")
                
        except Exception as conversion_error:
            # Try to clean up temp file even if conversion fails
            if os.path.exists(docx_path):
                try:
                    os.remove(docx_path)
                except:
                    pass
                    
            return jsonify({
                "status": "error", 
                "message": f"Failed to convert to PDF: {str(conversion_error)}. "
                           "Note: PDF conversion requires Microsoft Word or LibreOffice to be installed."
            }), 500
            
        # Check if PDF was created successfully
        if not os.path.exists(pdf_path):
            return jsonify({
                "status": "error", 
                "message": "Failed to create PDF. Please make sure Microsoft Word or LibreOffice is installed on your system."
            }), 500
        
        # Return the file info
        return jsonify({
            "status": "success",
            "message": f"PDF Resume successfully generated.",
            "filename": temp_pdf_filename,
            "file_path": pdf_path
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """Download the generated resume document and clean up temporary files"""
    try:
        # Send the file to the client for download
        response = send_file(filename, as_attachment=True)
        
        # Register a function to clean up the file after it's sent
        @response.call_on_close
        def cleanup():
            try:
                # Only delete files with temp_ prefix
                if filename.startswith('temp_') and os.path.exists(filename):
                    os.remove(filename)
                    print(f"Temporary file {filename} deleted after download")
            except Exception as e:
                print(f"Error cleaning up file: {str(e)}")
                
        return response
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/download-pdf/<filename>', methods=['GET'])
def download_pdf(filename):
    """Convert and download the generated resume as PDF"""
    try:
        # Get base name without extension
        base_name = os.path.splitext(filename)[0]
        pdf_filename = f"{base_name}.pdf"
        docx_path = os.path.join(os.getcwd(), filename)
        pdf_path = os.path.join(os.getcwd(), pdf_filename)
        
        # Convert DOCX to PDF using python-docx-pdf
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
        except ImportError:
            # Fallback method if docx2pdf not available
            return jsonify({"status": "error", "message": "PDF conversion requires docx2pdf module. Please install it with 'pip install docx2pdf'"}), 500
        
        # Check if PDF was created successfully
        if not os.path.exists(pdf_path):
            return jsonify({"status": "error", "message": "Failed to create PDF"}), 500
            
        return send_file(pdf_path, as_attachment=True)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/get_raw_json', methods=['POST'])
def get_raw_json():
    """Return the raw JSON data from Gemini API parsing"""
    try:
        # Get the parsed resume data from request
        resume_data = request.json
        
        if not resume_data:
            return jsonify({"status": "error", "message": "No data provided"}), 400
        
        # Return the raw JSON data
        return jsonify({
            "status": "success",
            "data": resume_data
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # thin line
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)

def add_work_experience(doc, job):
    # Create a 1-row, 2-column table
    table = doc.add_table(rows=2, cols=2)

    # Left cell (Company and Location)
    left_top = table.cell(0, 0)
    left_p = left_top.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_p.paragraph_format.space_after = Pt(0)  # Remove space after
    left_run_company = left_p.add_run(job['company'])
    left_run_company.italic = True
    left_run_company.bold = True

    left_bottom = table.cell(1, 0)
    left_p = left_bottom.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_p.paragraph_format.space_after = Pt(0)  # Remove space after
    left_run_role = left_p.add_run(job['role'])
    left_run_role.italic = True
    
    # Right cell (Position and Dates)
    right_top = table.cell(0, 1)
    right_p = right_top.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.space_after = Pt(0)  # Remove space after
    right_run_location = right_p.add_run(job['location'])
    right_run_location.bold = True
    right_run_location.italic = True

    right_bottom = table.cell(1, 1)
    right_p = right_bottom.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.space_after = Pt(0)  # Remove space after
    # Fix date ranges by explicitly replacing Unicode characters
    dates = job['dates'].replace('\u00e2', '-').replace('\u2013', '-').replace('\u2014', '-').replace('  ', ' ')
    right_run_dates = right_p.add_run(dates)
    right_run_dates.italic = True
    
    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_education(doc, education):
    # Create a 1-row, 2-column table
    table = doc.add_table(rows=2, cols=2)

    # Left cell (institution and degree)
    left_top = table.cell(0, 0)
    left_p = left_top.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_p.paragraph_format.space_after = Pt(0)  # Remove space after
    left_run_institution = left_p.add_run(education['institution'])
    left_run_institution.italic = True
    left_run_institution.bold = True

    left_bottom = table.cell(1, 0)
    left_bottom.width = Inches(5)  # Set width for left cell
    left_p = left_bottom.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_p.paragraph_format.space_after = Pt(0)  # Remove space after
    
    # Fix apostrophe character encoding issues
    degree_text = education['degree']
    degree_text = degree_text.replace('â', "'").replace('â', "'").replace('\u00e2', "'")
    
    left_run_degree = left_p.add_run(degree_text)
    left_run_degree.italic = True
    
    # Right cell (Location and Grad year)
    right_top = table.cell(0, 1)
    right_p = right_top.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.space_after = Pt(0)  # Remove space after
    right_run_location = right_p.add_run(education['location'])
    right_run_location.bold = True
    right_run_location.italic = True

    right_bottom = table.cell(1, 1)
    right_p = right_bottom.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.space_after = Pt(0)  # Remove space after
    dates = education['graduation_year'].replace('\u00e2', '-').replace('  ', ' ')
    right_run_dates = right_p.add_run(f"Graduation: {dates}")
    right_run_dates.italic = True

    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_section_separator(doc):
    doc.add_paragraph()  # Add an empty paragraph for spacing

def create_resume_document(resume_data, filename):
    if isinstance(resume_data, str):
        try:
            resume_data = json.loads(resume_data)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON string: {str(e)}") from e
        
    if 'contact' not in resume_data or not isinstance(resume_data['contact'], dict):
        raise ValueError("Missing or invalid contact information")
    
    # Create a new Document
    doc = Document()
    
    # ===== Set Default Styles =====
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Helvetica Neue'
    font.size = Pt(10)

    # Set 0.5 inch margins on all sides
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # ===== Contact Information =====
    # More compact header with name and contact info on same line
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    contact.add_run(resume_data['contact']['name']).bold = True
    
    # Contact info on same line with minimal spacing
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(0)  # No space after
    contact_p.paragraph_format.space_before = Pt(0)  # No space before
    
    # Add email with hyperlink
    email = resume_data['contact']['email'].strip()
    add_hyperlink(contact_p, email, f"mailto:{email}")
    contact_p.add_run(" || ").bold = True
    
    # Add phone number
    contact_p.add_run(resume_data['contact']['phone'].strip()).bold = True
    contact_p.add_run(" || ").bold = True
    
    # Add LinkedIn with hyperlink - keep exactly as provided in the input
    linkedin_text = resume_data['contact']['linkedin'].strip()
    linkedin_url = linkedin_text
    
    # Ensure URL has https:// prefix for proper linking if not already present
    if not linkedin_url.startswith("http"):
        linkedin_url = "https://" + linkedin_url
    
    add_hyperlink(contact_p, linkedin_text, linkedin_url)
    
    # Add title on same line or with minimal spacing
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(0)  # No space after
    title_p.paragraph_format.space_before = Pt(2)  # Minimal space before
    title_run = title_p.add_run(resume_data['title'])
    title_run.underline = True
    title_run.bold = True

    # We still need some minimal separation before content sections
    doc.add_paragraph("", style="Normal")
    
    # ===== Professional Summary =====
    summary_heading = doc.add_paragraph()
    summary_heading.add_run('Professional Summary:').bold = True

    summary = doc.add_paragraph(resume_data['summary'])

    # ===== Technical Skills Heading =====
    skills_heading = doc.add_paragraph()
    skills_heading.add_run('Technical Skills:').bold = True

    # Create table with 1 row (to define column widths)
    skills_table = doc.add_table(rows=1, cols=2)
    skills_table.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Align with document margins
    skills_table.autofit = False  # Prevent auto-adjustment
    set_table_borders(skills_table)

    # Set widths for each column
    skills_table.columns[0].width = Inches(2)
    skills_table.columns[1].width = Inches(5.5)

    # Remove the first placeholder row
    skills_table._tbl.remove(skills_table.rows[0]._tr)

    # Set header texts
    row = skills_table.add_row().cells
    row[0].text = "Category"
    row[1].text = "Skills/Tools"

    # Make both headers bold
    row[0].paragraphs[0].runs[0].bold = True
    row[1].paragraphs[0].runs[0].bold = True

    # Remove spacing after paragraph
    for cell in row:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = 0

    # Add actual data rows - preserve the exact order from input data
    if isinstance(resume_data['technical_skills'], dict):
        # We need to iterate through items in the exact order they appear in the original data
        for category, skills in resume_data['technical_skills'].items():
            row = skills_table.add_row().cells
            row[0].text = category
            row[1].text = ", ".join(skills) if isinstance(skills, list) else skills

            # Bold left column text
            if row[0].paragraphs[0].runs:
                row[0].paragraphs[0].runs[0].bold = True

            # Remove spacing after paragraph
            for cell in row:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = 0

    add_section_separator(doc)

    # ===== Work Experience =====
    work_exp_heading = doc.add_paragraph()
    work_exp_heading.add_run('Work Experience').bold = True

    for job in resume_data['work_experience']:
        add_work_experience(doc, job)
        
        # Achievements
        for achievement in job['achievements']:
            achievement = achievement.replace("●", "").strip()
            doc.add_paragraph(achievement, style='ListBullet')

    # ===== Education =====
    edu_heading = doc.add_paragraph()
    edu_heading.add_run('Education:').bold = True
    add_education(doc, resume_data['education'])

    # Save the document
    doc.save(filename)
    print(f"Resume saved as {filename}")
    return filename

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a run element
    new_run = OxmlElement('w:r')
    
    # Create a bold element if needed
    if paragraph.runs and paragraph.runs[-1].bold:
        b = OxmlElement('w:b')
        new_run.append(b)
    
    # Create a rPr element
    rPr = OxmlElement('w:rPr')
    
    # Add color if you want it
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue color
    rPr.append(color)
    
    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # Create a text element
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink

if __name__ == '__main__':
    app.run(debug=True)